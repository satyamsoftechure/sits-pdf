import json
import logging
import os
from multiprocessing import Pool, cpu_count
from time import perf_counter
from typing import AnyStr, IO, Union
import fitz
from docx import Document
from .page.Page import Page
from .page.Pages import Pages

v = list(map(int, fitz.VersionBind.split(".")))
if v < [1, 19, 0] or [1, 23, 8] < v < [1, 23, 16]:
    raise SystemExit(
        "1.19.0 <= PyMuPDF <= 1.23.8, or PyMuPDF>=1.23.16 is required for pdf2docx."
    )

logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")


class Converter:
    def __init__(
        self, pdf_file: str = None, password: str = None, stream: bytes = None
    ):
        self.filename_pdf = pdf_file
        self.password = str(password or "")

        if not pdf_file and not stream: 
            raise ValueError("Either pdf_file or stream must be given.")

        if stream:
            self._fitz_doc = fitz.Document(stream=stream)

        else:
            self._fitz_doc = fitz.Document(pdf_file)
        self._pages = Pages()

    @property
    def fitz_doc(self):
        return self._fitz_doc

    @property
    def pages(self):
        return self._pages

    def close(self):
        self._fitz_doc.close()

    @property
    def default_settings(self):
        return {
            "debug": False,
            "ocr": 0,
            "ignore_page_error": True,
            "multi_processing": False,
            "cpu_count": 0,
            "min_section_height": 20.0,
            "connected_border_tolerance": 0.5,
            "max_border_width": 6.0,
            "min_border_clearance": 2.0,
            "float_image_ignorable_gap": 5.0,
            "page_margin_factor_top": 0.5,
            "page_margin_factor_bottom": 0.5,
            "shape_min_dimension": 2.0,
            "max_line_spacing_ratio": 1.5,
            "line_overlap_threshold": 0.9,
            "line_break_width_ratio": 0.5,
            "line_break_free_space_ratio": 0.1,
            "line_separate_threshold": 5.0,
            "new_paragraph_free_space_ratio": 0.85,
            "lines_left_aligned_threshold": 1.0,
            "lines_right_aligned_threshold": 1.0,
            "lines_center_aligned_threshold": 2.0,
            "clip_image_res_ratio": 4.0,
            "min_svg_gap_dx": 15.0,
            "min_svg_gap_dy": 2.0,
            "min_svg_w": 2.0,
            "min_svg_h": 2.0,
            "extract_stream_table": False,
            "parse_lattice_table": True,
            "parse_stream_table": True,
            "delete_end_line_hyphen": False,
        }

    def parse(self, start: int = 0, end: int = None, pages: list = None, **kwargs):
        return (
            self.load_pages(start, end, pages)
            .parse_document(**kwargs)
            .parse_pages(**kwargs)
        )

    def load_pages(self, start: int = 0, end: int = None, pages: list = None):
        logging.info(self._color_output("[1/4] Opening document..."))
        if self._fitz_doc.needs_pass:
            if not self.password:
                raise ConversionException(f"Require password for {self.filename_pdf}.")

            elif not self._fitz_doc.authenticate(self.password):
                raise ConversionException("Incorrect password.")
        num = len(self._fitz_doc)
        self._pages.reset([Page(id=i, skip_parsing=True) for i in range(num)])
        page_indexes = self._page_indexes(start, end, pages, num)
        for i in page_indexes:
            self._pages[i].skip_parsing = False

        return self

    def parse_document(self, **kwargs):
        logging.info(self._color_output("[2/4] Analyzing document..."))

        self._pages.parse(self.fitz_doc, **kwargs)
        return self

    def parse_pages(self, **kwargs):
        logging.info(self._color_output("[3/4] Parsing pages..."))

        pages = [page for page in self._pages if not page.skip_parsing]
        num_pages = len(pages)
        for i, page in enumerate(pages, start=1):
            pid = page.id + 1
            logging.info("(%d/%d) Page %d", i, num_pages, pid)
            try:
                page.parse(**kwargs)
            except Exception as e:
                if not kwargs["debug"] and kwargs["ignore_page_error"]:
                    logging.error(
                        "Ignore page %d due to parsing page error: %s", pid, e
                    )
                else:
                    raise ConversionException(f"Error when parsing page {pid}: {e}")

        return self

    def make_docx(self, filename_or_stream=None, **kwargs):
        logging.info(self._color_output("[4/4] Creating pages..."))
        parsed_pages = list(filter(lambda page: page.finalized, self._pages))
        if not parsed_pages:
            raise ConversionException("No parsed pages. Please parse page first.")

        if not filename_or_stream:
            if self.filename_pdf:
                filename_or_stream = f'{self.filename_pdf[0:-len(".pdf")]}.docx'
                if os.path.exists(filename_or_stream):
                    os.remove(filename_or_stream)
            else:
                raise ConversionException(
                    "Please specify a docx file name or a file-like object to write."
                )
        docx_file = Document()
        num_pages = len(parsed_pages)
        for i, page in enumerate(parsed_pages, start=1):
            if not page.finalized:
                continue
            pid = page.id + 1
            logging.info("(%d/%d) Page %d", i, num_pages, pid)
            try:
                page.make_docx(docx_file)
            except Exception as e:
                if not kwargs["debug"] and kwargs["ignore_page_error"]:
                    logging.error("Ignore page %d due to making page error: %s", pid, e)
                else:
                    raise MakedocxException(f"Error when make page {pid}: {e}")
        docx_file.save(filename_or_stream)

    def store(self):
        return {
            "filename": os.path.basename(self.filename_pdf),
            "page_cnt": len(self._pages),
            "pages": [page.store() for page in self._pages if page.finalized],
        }

    def restore(self, data: dict):
        if not self._pages:
            num = data.get("page_cnt", 100)
            self._pages.reset([Page(id=i, skip_parsing=True) for i in range(num)])

        for raw_page in data.get("pages", []):
            idx = raw_page.get("id", -1)
            self._pages[idx].restore(raw_page)

    def serialize(self, filename: str):
        with open(filename, "w", encoding="utf-8") as f:
            f.write(json.dumps(self.store(), indent=4))

    def deserialize(self, filename: str):
        with open(filename, "r") as f:
            data = json.load(f)
        self.restore(data)

    def debug_page(
        self,
        i: int,
        docx_filename: str = None,
        debug_pdf: str = None,
        layout_file: str = None,
        **kwargs,
    ):
        path, filename = os.path.split(self.filename_pdf)
        if not debug_pdf:
            debug_pdf = os.path.join(path, f"debug_{filename}")
        if not layout_file:
            layout_file = os.path.join(path, "..layout.json")
        kwargs.update(
            {"debug": True, "debug_doc": fitz.Document(), "debug_filename": debug_pdf}
        )

        self.convert(docx_filename, pages=[i], **kwargs)
        self.serialize(layout_file)

    def convert(
        self,
        docx_filename: Union[str, IO[AnyStr]] = None,
        start: int = 0,
        end: int = None,
        pages: list = None,
        **kwargs,
    ):
        t0 = perf_counter()
        logging.info("Start to convert %s", self.filename_pdf)
        settings = self.default_settings
        settings.update(kwargs)

        if pages and settings["multi_processing"]:
            raise ConversionException(
                "Multi-processing works for continuous pages "
                'specified by "start" and "end" only.'
            )

        if settings["multi_processing"]:
            self._convert_with_multi_processing(docx_filename, start, end, **settings)
        else:
            self.parse(start, end, pages, **settings).make_docx(
                docx_filename, **settings
            )

        logging.info("Terminated in %.2fs.", perf_counter() - t0)

    def extract_tables(
        self, start: int = 0, end: int = None, pages: list = None, **kwargs
    ):
        settings = self.default_settings
        settings.update(kwargs)
        self.parse(start, end, pages, **settings)

        tables = []
        for page in self._pages:
            if page.finalized:
                tables.extend(page.extract_tables(**settings))

        return tables

    def _convert_with_multi_processing(
        self, docx_filename: str, start: int, end: int, **kwargs
    ):
        cpu = (
            min(kwargs["cpu_count"], cpu_count())
            if kwargs["cpu_count"]
            else cpu_count()
        )
        prefix = "pages"
        vectors = [
            (
                i,
                cpu,
                start,
                end,
                self.filename_pdf,
                self.password,
                kwargs,
                f"{prefix}-{i}.json",
            )
            for i in range(cpu)
        ]
        pool = Pool()
        pool.map(self._parse_pages_per_cpu, vectors, 1)
        for i in range(cpu):
            filename = f"{prefix}-{i}.json"
            if not os.path.exists(filename):
                continue
            self.deserialize(filename)
            os.remove(filename)
        self.make_docx(docx_filename, **kwargs)

    @staticmethod
    def _parse_pages_per_cpu(vector):
        idx, cpu, s, e, pdf_filename, password, kwargs, json_filename = vector
        cv = Converter(pdf_filename, password)
        cv.load_pages()
        e = e or len(cv.fitz_doc)
        all_indexes = range(s, e)
        num_pages = len(all_indexes)
        m = int(num_pages / cpu)
        n = num_pages % cpu
        seg_size = m + int(idx < n)
        seg_from = (m + 1) * idx + min(n - idx, 0)
        seg_to = min(seg_from + seg_size, num_pages)
        page_indexes = [all_indexes[i] for i in range(seg_from, seg_to)]
        for page in cv.pages:
            page.skip_parsing = True
        for i in page_indexes:
            cv.pages[i].skip_parsing = False
        cv.parse_document(**kwargs).parse_pages(**kwargs).serialize(json_filename)
        cv.close()

    @staticmethod
    def _page_indexes(start, end, pages, pdf_len):
        if pages:
            indexes = [int(x) for x in pages]
        else:
            end = end or pdf_len
            s = slice(int(start), int(end))
            indexes = range(pdf_len)[s]

        return indexes

    @staticmethod
    def _color_output(msg):
        return f"\033[1;36m{msg}\033[0m"


class ConversionException(Exception):
    pass


class MakedocxException(ConversionException):
    pass
